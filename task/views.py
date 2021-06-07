from django.shortcuts import render,redirect
from django.core.files import File
from docx import Document
from docx.shared import Inches
from task.models import FileModel
from Assignment import settings
import textract
# open connection to Word Document

def index(request):
    if request.method == "POST":
        # collect all input values from client
        content = request.POST['content']
        filename = request.POST['filename']
        # created object for Document()
        document = Document()
        content = document.add_paragraph(content)
        file = document.save(filename+'.docx')
        obj = FileModel(word_file=filename+'.docx')
        obj.save()
        return redirect('index')
    all_files = FileModel.objects.all()
    return render(request,'index.html',{'all_files':all_files})

def edit(request,id):
    data ={}
    if request.method == "POST":
        content = request.POST["content"]
        obj=FileModel.objects.get(id=id)
        document = Document()
        content = document.add_paragraph(content)
        file = document.save(str(obj.word_file))
        obj.save()
        return redirect('index')
    obj=FileModel.objects.get(id=id)
    text = textract.process(str(obj.word_file))
    text = text.decode("utf-8")
    return render(request,'edit.html',{"data":text,"id":id})
